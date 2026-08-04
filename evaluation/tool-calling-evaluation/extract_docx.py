from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


SOURCE = Path("/Users/zhanhuilin/Downloads/Analytica 工具调用能力评测.docx")
OUTPUT = Path("evaluation/tool-calling-evaluation/docx-structure.json")


def iter_blocks(document: DocumentType):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


document = Document(SOURCE)
blocks: list[dict[str, object]] = []
for index, block in enumerate(iter_blocks(document), start=1):
    if isinstance(block, Paragraph):
        blocks.append(
            {
                "index": index,
                "type": "paragraph",
                "style": block.style.name if block.style else None,
                "text": block.text,
                "runs": [run.text for run in block.runs],
            }
        )
    else:
        blocks.append(
            {
                "index": index,
                "type": "table",
                "rows": [
                    ["\n".join(paragraph.text for paragraph in cell.paragraphs) for cell in row.cells]
                    for row in block.rows
                ],
            }
        )

OUTPUT.write_text(
    json.dumps(
        {
            "source": str(SOURCE),
            "paragraphCount": len(document.paragraphs),
            "tableCount": len(document.tables),
            "blocks": blocks,
        },
        ensure_ascii=False,
        indent=2,
    )
    + "\n",
    encoding="utf-8",
)
print(f"extracted {len(blocks)} blocks")
